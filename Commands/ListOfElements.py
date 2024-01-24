import os
import os
import docx
import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt
from docxtpl import DocxTemplate

import Constants
import DataModel
import DataParser
import Editors

def export_to_word(dm: DataModel):
    """
    Экспортирует полученный список в формат Word
    """
    # if dm.TemplatesPath != "":
    #     path_to_template = dm.TemplatesPath
    # else:
    #     path_to_template = obshiy_perechen.prog_dir + "\\Шаблоны"
    #
    # if not civilian:
    #     doc = Document(path_to_template + '\\Шаблон ПЭ.docx')
    # else:
    #     doc = Document(path_to_template + '\\Шаблон ПЭ Гражданский.docx')
    #doc = Document('C:\\Users\\EDN\\Desktop\\mainDocsUI\\Шаблоны\\Шаблон ПЭ.docx')

    doc = Document(dm.TemplatesPath + '\\Шаблон ПЭ.docx')
    style = doc.styles['Normal']
    font = style.font
    font.name = 'T-FLEX Type A'
    font.size = Pt(12)

    # Получаем итераторы таблиц в документе и строк в нем
    tables = iter(doc.tables)
    table = next(tables)
    rows = iter(table.rows)

    next(rows)
    row = next(rows)
    row_index = 1



    try:
        for char in sorted(dm.dict_chars.keys()):
            for component_index, component in enumerate(dm.dict_chars[char]):
                component.split_name(shift_threshold=56)
                component.split_desig(shift_treshold=54)

                # Вставка наименования категории компонентов
                if component_index == 0:
                    try:
                        row = next(rows)
                        row_index += 1
                    except StopIteration:
                        table = next(tables)
                        rows = iter(table.rows)
                        next(rows)
                        row = next(rows)

                        row_index = 1

                    rows_len = len(table.rows) - 1
                    if row_index + len(component.name) > rows_len or row_index + len(component.splt_desig) > rows_len \
                            or (isinstance(component.quantity, list) and row_index + len(component.quantity) > rows_len) or \
                            (isinstance(component.man, list) and row_index + len(component.man) > rows_len):
                        table = next(tables)
                        rows = iter(table.rows)
                        next(rows)
                        row = next(rows)

                        row_index = 1

                    cat_name = ''
                    if len(dm.dict_chars[char]) > 1:
                        for desig, d_cat_name in Constants.cat_names_plural.items():
                            if char == desig:
                                cat_name = d_cat_name

                                if char in dm.components_one_manuf_categories.keys():
                                    cat_name += f', {dm.components_one_manuf_categories[char]}'
                                break

                        row.cells[1].text = cat_name
                        row.cells[1].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                        row.cells[1].paragraphs[0].runs[0].underline = docx.enum.text.WD_UNDERLINE.SINGLE

                        try:
                            row = next(rows)
                            row_index += 1
                        except StopIteration:
                            table = next(tables)
                            rows = iter(table.rows)
                            next(rows)
                            row = next(rows)

                            row_index = 1
                    # Если компонент один
                    else:
                        for desig, d_cat_name in Constants.cat_names_singular.items():
                            if char == desig:
                                cat_name = d_cat_name

                        component.split_name(shift_threshold=56, cat_name=cat_name)

                # Вставка строки в документ и ее оформление
                rows_len = len(table.rows) - 1
                if row_index + len(component.name) - 1 > rows_len or row_index + len(component.splt_desig) - 1 > rows_len \
                        or (
                        isinstance(component.quantity, list) and row_index + len(component.quantity) - 1 > rows_len) or \
                        (isinstance(component.man, list) and row_index + len(component.man) - 1 > rows_len):
                    table = next(tables)
                    rows = iter(table.rows)
                    next(rows)
                    row = next(rows)

                    row_index = 1

                row.cells[2].text = str(component.quantity)
                row.cells[2].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

                row.cells[3].text = component.comment
                row.cells[3].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

                while True:
                    try:
                        row.cells[0].text = str(component.splt_desig[0])
                        row.cells[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

                        del component.splt_desig[0]
                    except IndexError:
                        pass

                    try:
                        row.cells[1].text = str(component.name[0])
                        row.cells[1].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

                        del component.name[0]
                    except IndexError:
                        pass

                    if (char == 'C' or char == 'R') and 'ТУ' not in component.man:
                        try:
                            row.cells[3].text = str(component.manpnb)
                            row.cells[3].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

                            del component.name[0]
                        except IndexError:
                            pass

                    if component.name != [] or component.splt_desig != []:
                        try:
                            row = next(rows)
                            row_index += 1
                        except StopIteration:
                            table = next(tables)
                            rows = iter(table.rows)
                            next(rows)
                            row = next(rows)
                            start_row = True

                            row_index = 1
                    else:
                        break

                try:
                    row = next(rows)
                    row_index += 1
                except StopIteration:
                    table = next(tables)
                    rows = iter(table.rows)
                    next(rows)
                    row = next(rows)
                    start_row = True

                    row_index = 1
    except StopIteration:
        pass

    save_dir1 = dm.files[0][:-len(dm.files[0].split("\\")[-1])] + "Перечни\\"
    save_dir = Editors.only_dir(dm.files[0]) + "Перечни\\"
    new_name = pd.read_excel(dm.files[0], sheet_name='BOM', header=None).loc[7, 10].split(' ')[0]
    deviceName = pd.read_excel(dm.files[0], sheet_name='BOM', header=None).loc[2, 0]
    if new_name == '':
        print("Не заполнено поле 'Первичная Применямость'!")
    if deviceName == '':
        print("Не заполнено поле 'Наименование1'!")

    if not os.path.exists(save_dir):
        os.makedirs(save_dir)
    doc.save(save_dir + new_name + ' ПЭ (без полей).docx')

    doc = DocxTemplate(save_dir + new_name + ' ПЭ (без полей).docx')

    doc.render(dm.data_for_first_page)

    # if civilian:
    #     new_name += ' гражданский'
    doc.save(save_dir + new_name + ' ПЭ.docx')
    print(f"Файл сохранен по пути: {save_dir + new_name + '.docx'}")


def Execute(dm: DataModel):
    for i in range(len(dm.files)):

        print(f"\nФормирую {i+1}-й ПЭ\n"
              f"=============")


        DataParser.test = True

        DataParser.no_perechen = 0
        print("Получаю данные из перечней элементов...")
        DataParser.get_dfs(dm, i)
        print("Формирую общую таблицу элементов...")
        DataParser.get_components(dm)
        print("Проверяю примечения на регулирование...")
        DataParser.split_to_adjustable(dm)
        print("Вставляю готовый перечень в шаблон...")
        export_to_word(dm)

    print("\n=============\n"
          "Готово!\n\n\n")
